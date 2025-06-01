from pathlib import Path
import io
from typing import Tuple
from config.log_config import get_logger
import docx
from zipfile import ZipFile
from itertools import groupby
from pdfminer.high_level import extract_text_to_fp
from pdfminer.layout import LAParams


logger = get_logger()


def get_file_extension(file_path: str) -> str:
    """Get the lowercase file extension without the dot."""
    return Path(file_path).suffix.lower()[1:]


def extract_text_from_pdf(pdf_file: io.BytesIO) -> str:
    """
    Extract text from a PDF file stored as a BytesIO object,
    preserving formatting and spacing as much as possible.
    Returns a single string without newline characters.

    :param pdf_file: BytesIO object containing the PDF data
    :return: Extracted text as a single string, with formatting preserved but newlines removed
    """
    output_string = io.StringIO()
    laparams = LAParams(
        line_margin=0.5, word_margin=0.1, char_margin=2.0, all_texts=True
    )

    extract_text_to_fp(
        pdf_file, output_string, laparams=laparams, output_type="text", codec="utf-8"
    )

    text = output_string.getvalue()

    # Post-processing to improve readability
    lines = text.split("\n")
    processed_lines = []
    for line in lines:
        # Remove lines that are just whitespace
        if line.strip():
            # Ensure bullet points and other list markers are followed by a space
            line = line.replace("•", "• ").replace("◦", "◦ ")
            processed_lines.append(line)

    # Join lines, preserving paragraph structure
    processed_text = "\n".join(processed_lines)

    # Remove multiple consecutive blank lines
    processed_text = "\n".join(line for line, _ in groupby(processed_text.split("\n")))

    # Remove all newline characters to create a single string
    final_text = processed_text.replace("\n", " ").strip()

    return final_text


def extract_text_from_docx(file: io.BytesIO) -> str:
    """
    Extract text from a DOCX file, including tables.

    Args:
        file: BytesIO object containing the DOCX file

    Returns:
        str: Extracted text content
    """
    try:
        doc = docx.Document(file)
        text_content = []

        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():  # Skip empty paragraphs
                text_content.append(para.text)

        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():  # Skip empty cells
                        row_text.append(cell.text)
                if row_text:  # Skip empty rows
                    text_content.append(" | ".join(row_text))

        return " ".join(text_content)

    except Exception as e:
        # Fallback for corrupted DOCX files
        logger.error("Could not extract text from docx", e)
        return None


def extract_text_from_odt(file: io.BytesIO) -> str:
    """
    Extract text from an ODT file.

    Args:
        file: BytesIO object containing the ODT file

    Returns:
        str: Extracted text content
    """
    with ZipFile(file) as odt_file:
        content = odt_file.read("content.xml").decode()
        # Basic XML parsing to extract text
        import re
        from html import unescape

        text = re.sub(r"<[^>]+>", " ", content)
        text = unescape(text)
        return " ".join(text.split())


def extract_text_from_document(file: io.BytesIO, file_path: str) -> Tuple[str, str]:
    """
    Process a document file and return its text content and actual file type.

    Args:
        file: BytesIO object containing the file
        file_path: Original file path for type detection

    Returns:
        Tuple[str, str]: (extracted text, actual file type used)

    Raises:
        ValueError: If file type is unsupported
    """
    extension = get_file_extension(file_path)

    try:
        if extension == "pdf":
            return extract_text_from_pdf(file), "pdf"

        elif extension == "doc":
            raise ValueError(
                "Old .doc format is not supported. Please convert to .docx or pdf first."
            )

        elif extension == "docx":
            return extract_text_from_docx(file), "docx"

        elif extension == "odt":
            return extract_text_from_odt(file), "odt"

        elif extension == "txt":
            return file.getvalue().decode("utf-8", errors="ignore"), "txt"

        else:
            raise ValueError(f"Unsupported file type: {extension}")

    except Exception as e:
        raise ValueError(f"Error processing {extension} file: {str(e)}")
