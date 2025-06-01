from pathlib import Path
import io
from typing import Tuple
from config.log_config import get_logger
import logging

# Import with fallbacks to handle potential missing dependencies
try:
    import fitz  # PyMuPDF
    _HAS_PYMUPDF = True
except ImportError:
    _HAS_PYMUPDF = False
    logging.getLogger(__name__).warning("PyMuPDF not available - PDF processing will be limited")

try:
    import docx
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False
    logging.getLogger(__name__).warning("python-docx not available - DOCX processing will be disabled")

from itertools import groupby


logger = get_logger()


def get_file_extension(file_path: str) -> str:
    """Get the lowercase file extension without the dot."""
    return Path(file_path).suffix.lower()[1:]


def extract_text_from_pdf(pdf_file: io.BytesIO) -> str:
    """
    Extract text from a PDF file stored as a BytesIO object,
    preserving formatting and spacing as much as possible.
    
    Args:
        pdf_file: BytesIO object containing the PDF data
    
    Returns:
        Extracted text as a string
    
    Raises:
        ValueError: If PDF extraction fails or PyMuPDF is not available
    """
    if not _HAS_PYMUPDF:
        raise ValueError("PDF extraction requires PyMuPDF (fitz) which is not installed")
    
    try:
        doc = fitz.open(stream=pdf_file, filetype="pdf")
        text_parts = []
        
        for page in doc:
            text_parts.append(page.get_text())
            
        # Join all text parts with proper spacing
        text = "\n".join(text_parts)
        
        # Post-processing to improve readability
        lines = text.split("\n")
        processed_lines = []
        for line in lines:
            # Remove lines that are just whitespace
            if line.strip():
                # Ensure bullet points and other list markers are followed by a space
                line = line.replace("•", "• ").replace("◦", "◦ ")
                processed_lines.append(line)

        # Join lines
        processed_text = "\n".join(processed_lines)
        
        # Remove multiple consecutive blank lines
        processed_text = "\n".join(line for line, _ in groupby(processed_text.split("\n")))
        
        return processed_text
    except Exception as e:
        error_msg = f"Error extracting text from PDF: {str(e)}"
        logger.error(error_msg, exc_info=True)
        raise ValueError(error_msg)


def extract_text_from_docx(file: io.BytesIO) -> str:
    """
    Extract text from a DOCX file, including tables.

    Args:
        file: BytesIO object containing the DOCX file

    Returns:
        str: Extracted text content
        
    Raises:
        ValueError: If DOCX extraction fails or python-docx is not available
    """
    if not _HAS_DOCX:
        raise ValueError("DOCX extraction requires python-docx which is not installed")
    
    try:
        doc = docx.Document(file)
        text_content = []

        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():  # Skip empty paragraphs
                text_content.append(para.text)

        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():  # Skip empty cells
                        row_text.append(cell.text)
                if row_text:  # Skip empty rows
                    text_content.append(" | ".join(row_text))

        return "\n".join(text_content)

    except Exception as e:
        error_msg = f"Error extracting text from DOCX: {str(e)}"
        logger.error(error_msg, exc_info=True)
        raise ValueError(error_msg)


def extract_text_from_document(file: io.BytesIO, file_path: str) -> Tuple[str, str]:
    """
    Process a document file and return its text content and actual file type.

    Args:
        file: BytesIO object containing the file
        file_path: Original file path for type detection

    Returns:
        Tuple[str, str]: (extracted text, actual file type used)

    Raises:
        ValueError: If file type is unsupported or text extraction fails
    """
    extension = Path(file_path).suffix.lower()
    logger.info(f"Extracting text from file with extension: {extension}")

    try:
        if extension == ".pdf":
            return extract_text_from_pdf(file), "pdf"

        elif extension == ".doc":
            raise ValueError(
                "Old .doc format is not supported. Please convert to .docx or pdf first."
            )

        elif extension == ".docx":
            return extract_text_from_docx(file), "docx"

        elif extension == ".txt":
            return file.getvalue().decode("utf-8", errors="ignore"), "txt"

        else:
            raise ValueError(f"Unsupported file type: {extension}")

    except Exception as e:
        raise ValueError(f"Error processing {extension} file: {str(e)}") 